"""DOCX 文档解析器。

使用 python-docx 提取段落、标题层级、表格。
"""

from __future__ import annotations

import hashlib
import os
from pathlib import Path

from src.domain import (
    Document, ParseQuality, Section, SectionType,
)
from src.loaders.base import BaseLoader


def _heading_level_from_style(style_name: str | None) -> int | None:
    """从 Word 样式名推断标题级别。

    Word 内置标题样式: "Heading 1" ~ "Heading 6"
    中文版: "标题 1" ~ "标题 6"
    """
    if not style_name:
        return None
    style_lower = style_name.lower().strip()

    # 英文版 Word
    if style_lower.startswith("heading "):
        try:
            level = int(style_lower.split()[-1])
            if 1 <= level <= 6:
                return level
        except ValueError:
            pass

    # 中文版 Word
    if style_lower.startswith("标题 "):
        try:
            level = int(style_lower.split()[-1])
            if 1 <= level <= 6:
                return level
        except ValueError:
            pass

    return None


class DocxLoader(BaseLoader):
    """DOCX 文档解析器。"""

    SUPPORTED_EXTENSIONS = [".docx"]

    def load(self, filepath: str) -> Document:
        """解析 DOCX 文件，返回 Document 对象。

        提取段落（含标题层级）和表格，构建 Section 列表。
        """
        filepath = os.path.realpath(os.path.abspath(os.path.expanduser(filepath)))
        stat = os.stat(filepath)
        source_id = hashlib.sha256(
            os.path.normcase(filepath).encode("utf-8")
        ).hexdigest()
        content_sha256 = self._sha256_file(filepath)

        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(filepath)
        except Exception as e:
            raise ValueError(f"无法解析 DOCX 文件 {filepath}: {e}") from e

        sections: list[Section] = []
        char_offset = 0
        headings_stack: list[tuple[int, str]] = []

        # 提取段落
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            style_name = para.style.name if para.style else None
            heading_level = _heading_level_from_style(style_name)

            # 更新标题栈
            if heading_level is not None:
                while headings_stack and headings_stack[-1][0] >= heading_level:
                    headings_stack.pop()
                headings_stack.append((heading_level, text))

            heading_path = ""
            if headings_stack:
                heading_path = " > ".join(h[1] for h in headings_stack)

            section_type = SectionType.HEADING if heading_level is not None else SectionType.PARAGRAPH
            sections.append(Section(
                section_type=section_type,
                text=text,
                heading_level=heading_level,
                heading_path=heading_path,
                char_start=char_offset,
                char_end=char_offset + len(text),
            ))
            char_offset += len(text) + 1

        # 提取表格
        for table_idx, table in enumerate(doc.tables):
            table_text = self._extract_table_text(table)
            if not table_text.strip():
                continue

            sections.append(Section(
                section_type=SectionType.TABLE,
                text=table_text,
                heading_path=" > ".join(h[1] for h in headings_stack) if headings_stack else "",
                char_start=char_offset,
                char_end=char_offset + len(table_text),
                metadata={"table_index": table_idx, "rows": len(table.rows), "cols": len(table.columns)},
            ))
            char_offset += len(table_text) + 1

        return Document(
            source_id=source_id,
            source_path=filepath,
            source_name=os.path.basename(filepath),
            file_type="docx",
            sections=sections,
            parse_quality=ParseQuality.STRUCTURED,
            parser_version="1.0",
            content_sha256=content_sha256,
            source_size=stat.st_size,
            source_mtime_ns=stat.st_mtime_ns,
        )

    def _extract_table_text(self, table) -> str:
        """将表格转换为文本表示。"""
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        return "\n".join(rows)

    @staticmethod
    def _sha256_file(filepath: str) -> str:
        digest = hashlib.sha256()
        with open(filepath, "rb") as stream:
            for block in iter(lambda: stream.read(1024 * 1024), b""):
                digest.update(block)
        return digest.hexdigest()
