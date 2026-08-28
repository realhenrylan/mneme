"""2.3 解析失败可见性端到端验收矩阵（TDD）。

验收口径（plans/STAGE2-PART2-DESIGN-2026-08-28.md Part 2 / F3）：
「目标格式的解析失败可见、可回退」——以程序化构造的真实 PDF/DOCX/TXT
夹具走完整 ``_load_index_chunks`` 生产路径，断言三条可见性通道同时成立：

1. CLI stdout：低质量/降级/零块警告文本可见；
2. 诊断 sink：结构化条目（parse_quality / is_low_quality / chunk_count /
   parse_degraded / error）；
3. TUI 过滤器（``tui.diagnostics.parse_diagnostics_warnings``）命中问题条目
   且对正常文件静默。

矩阵：原生 PDF / 仿真扫描 PDF（fitz 空白页）/ 空文本 txt / loader 异常 /
正常 docx 与 txt。全矩阵通过 = ``STAGE2_23_ACCEPTED``。
"""
from __future__ import annotations

import fitz
import pytest

import src.rag as rag
from tui.diagnostics import parse_diagnostics_warnings


# ── 夹具生成 ─────────────────────────────────────────────────────

def _native_pdf(tmp_path):
    """含真实文本层的 PDF（fitz 程序化生成，无二进制 fixture）。"""
    path = tmp_path / "native.pdf"
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "1. 章节标题 城市概况")
    page.insert_text((72, 100), "正文段落内容 sufficient text " * 6)
    doc.save(str(path))
    doc.close()
    return path


def _scanned_pdf(tmp_path):
    """仿真扫描件：三页全部无文本层（fitz 空白页）。"""
    path = tmp_path / "scanned.pdf"
    doc = fitz.open()
    for _ in range(3):
        doc.new_page()
    doc.save(str(path))
    doc.close()
    return path


def _docx(tmp_path):
    from docx import Document as DocxDocument
    path = tmp_path / "doc.docx"
    document = DocxDocument()
    document.add_paragraph("第一节 标题")
    document.add_paragraph("正文内容，用于结构化分块验证。")
    document.save(str(path))
    return path


def _run_chunks(path, sink):
    return rag._load_index_chunks(str(path), diagnostics_sink=sink)


# ── 矩阵 ─────────────────────────────────────────────────────────

class TestNativePdf:
    def test_clean_parse_all_channels_quiet(self, tmp_path, capsys):
        sink: list = []
        chunks, _, _, _, _, _ = _run_chunks(_native_pdf(tmp_path), sink)
        assert chunks
        assert sink[0]["parse_quality"] in ("native_text", "structured")
        assert sink[0]["is_low_quality"] is False
        assert sink[0]["chunk_count"] > 0
        assert parse_diagnostics_warnings({"parse_diagnostics": sink}) == []
        assert "低质量" not in capsys.readouterr().out


class TestScannedPdf:
    def test_low_quality_visible_on_all_channels(self, tmp_path, capsys):
        sink: list = []
        chunks, _, _, _, _, _ = _run_chunks(_scanned_pdf(tmp_path), sink)
        out = capsys.readouterr().out
        # CLI 通道：低质量 + 零块双警告
        assert "低质量解析" in out
        assert "解析产出 0 块" in out
        # sink 通道：结构化标记
        assert chunks == []
        assert sink[0]["is_low_quality"] is True
        assert sink[0]["chunk_count"] == 0
        # TUI 通道：过滤器命中
        assert parse_diagnostics_warnings({"parse_diagnostics": sink})


class TestEmptyText:
    def test_zero_chunk_visible_not_silent(self, tmp_path, capsys):
        path = tmp_path / "empty.txt"
        path.write_text("", encoding="utf-8")
        sink: list = []
        chunks, _, _, _, _, _ = _run_chunks(path, sink)
        assert chunks == []
        assert sink[0]["chunk_count"] == 0
        assert "解析产出 0 块" in capsys.readouterr().out
        assert parse_diagnostics_warnings({"parse_diagnostics": sink})


class TestLoaderFailureDegraded:
    def test_degraded_visible_and_flagged(self, tmp_path, monkeypatch,
                                          capsys):
        from src.loaders import LoaderRegistry
        # 夹具须为旧路径可解析的格式（txt 纯读）——降级路径本身不是
        # 无条件兜底：旧路径也无法解析的文件会整体失败（另行记录）。
        path = tmp_path / "doc.txt"
        path.write_text("第一段内容。\n第二段内容。\n", encoding="utf-8")

        def _boom(self, filepath):
            raise RuntimeError("zip 损坏仿真")

        monkeypatch.setattr(LoaderRegistry, "load", _boom)
        sink: list = []
        chunks, _, _, _, _, source = _run_chunks(path, sink)
        out = capsys.readouterr().out
        assert chunks, "降级路径必须回退建成索引（可回退口径）"
        assert "[降级]" in out and "RuntimeError" in out
        assert "Traceback" not in out
        assert source["parse_degraded"] is True
        assert sink[0]["parse_degraded"] is True
        assert "RuntimeError" in sink[0]["error"]
        assert parse_diagnostics_warnings({"parse_diagnostics": sink})


class TestNormalDocxAndTxt:
    def test_clean_files_emit_no_noise(self, tmp_path, capsys):
        sink: list = []
        for path in (_docx(tmp_path), _write_txt(tmp_path)):
            before = len(sink)
            chunks, _, _, _, _, _ = _run_chunks(path, sink)
            assert chunks
        assert len(sink) == 2
        assert all(d["parse_degraded"] is False and d["chunk_count"] > 0
                   for d in sink)
        assert parse_diagnostics_warnings({"parse_diagnostics": sink}) == []
        assert "降级" not in capsys.readouterr().out


def _write_txt(tmp_path):
    path = tmp_path / "doc.txt"
    path.write_text("第一段内容。\n第二段内容。\n", encoding="utf-8")
    return path
