"""2.1 标准文档模型可追溯性的端到端验收测试（TDD）。

完成标准（路线图 2.1）：section、page、type、parser version 均可追溯——
即四项信息不仅存在于内存 Document 模型，还必须**持久化**到索引层：
chunk metadata（chroma 可查）与 manifest source record（账本可查）。

现状审计结论（2026-08-28）：section（section_heading/section_type）、
page、type（section_type/chunk_type）已落 metadata；parser_version 仅存于
Document 对象，索引层不可见 → 本测试锁定该缺口（RED），实现后转 GREEN。
"""
from __future__ import annotations

import fitz
import pytest

import src.rag as rag


def _native_pdf(tmp_path):
    path = tmp_path / "native.pdf"
    doc = fitz.open()
    page = doc.new_page()
    # 标题行用大字号（≥ _MIN_HEADING_FONT_SIZE 且与正文拉开差距），
    # 触发 pdf_loader 的字号标题检测——让 heading 路径端到端可证。
    page.insert_text((72, 72), "1. 章节标题 城市概况", fontsize=18)
    page.insert_text((72, 100), "正文段落内容 sufficient text " * 6,
                     fontsize=11)
    doc.save(str(path))
    doc.close()
    return path


def _docx(tmp_path):
    from docx import Document as DocxDocument
    path = tmp_path / "doc.docx"
    document = DocxDocument()
    document.add_heading("第一节 标题", level=2)
    document.add_paragraph("正文内容，用于结构化分块验证。")
    document.save(str(path))
    return path


def _txt(tmp_path):
    path = tmp_path / "doc.txt"
    path.write_text("第一段内容。\n第二段内容。\n", encoding="utf-8")
    return path


class TestChunkMetadataTraceability:
    """四项追溯字段必须出现在每个 chunk 的持久化 metadata 中。"""

    def test_pdf_chunks_carry_all_four_fields(self, tmp_path):
        chunks, metadatas, ids, _, _, _ = rag._load_index_chunks(
            str(_native_pdf(tmp_path)))
        assert metadatas
        for meta in metadatas:
            # 可追溯口径：字段在场且 type 非空；无标题节 heading 合法为空
            # （归属信息由 section_type + page 承载），不要求全非空。
            assert "section_heading" in meta, "section 必须可追溯"
            assert meta["section_type"], "type（section 级）必须可追溯"
            assert meta["parser_version"], "parser version 必须可追溯"
        # 标题路径端到端可证：至少一个块带非空 heading
        assert any(m["section_heading"] for m in metadatas)
        # page 仅 PDF 逐块可追溯：至少一个块带 page
        assert any("page" in m for m in metadatas), "page 必须可追溯"

    def test_docx_chunks_carry_required_fields(self, tmp_path):
        chunks, metadatas, ids, _, _, _ = rag._load_index_chunks(
            str(_docx(tmp_path)))
        assert metadatas
        for meta in metadatas:
            assert meta["section_heading"]
            assert meta["section_type"]
            assert meta["parser_version"] == "1.0"

    def test_txt_chunks_carry_required_fields(self, tmp_path):
        # 纯文本无标题结构：heading 合法为空，可追溯口径 = 字段在场 +
        # section_type 非空（paragraph）。
        chunks, metadatas, ids, _, _, _ = rag._load_index_chunks(
            str(_txt(tmp_path)))
        assert metadatas
        for meta in metadatas:
            assert "section_heading" in meta
            assert meta["section_type"]
            assert meta["parser_version"] == "1.0"

    def test_pdf_parser_version_is_2_0(self, tmp_path):
        chunks, metadatas, ids, _, _, _ = rag._load_index_chunks(
            str(_native_pdf(tmp_path)))
        assert all(m["parser_version"] == "2.0" for m in metadatas)


class TestManifestSourceRecordTraceability:
    """manifest source record 必须携带 parser version（账本级追溯）。"""

    def test_source_record_carries_parser_version(self, tmp_path):
        for path in (_native_pdf(tmp_path), _docx(tmp_path), _txt(tmp_path)):
            _, _, _, _, _, source = rag._load_index_chunks(str(path))
            assert source["parser_version"], f"{path.name} 缺 parser_version"

    def test_degraded_source_record_still_carry_parser_version(
            self, tmp_path, monkeypatch):
        # 降级到旧路径时 source record 由 _source_metadata 构建——
        # 旧路径产出的账本同样需要 parser 版本可追溯（统一 schema 口径）。
        from src.loaders import LoaderRegistry
        path = _txt(tmp_path)

        def _boom(self, filepath):
            raise RuntimeError("降级仿真")

        monkeypatch.setattr(LoaderRegistry, "load", _boom)
        _, _, _, _, _, source = rag._load_index_chunks(str(path))
        assert source["parse_degraded"] is True
        assert source["parser_version"], "降级路径也必须可追溯 parser version"
